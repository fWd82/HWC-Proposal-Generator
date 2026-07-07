from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "templates" / "huawei-commercial-proposal-template.docx"
RED = "C7002B"
DARK = "17212B"
MUTED = "68737D"
LIGHT = "F3F5F7"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=110, bottom=100, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_field(paragraph, instruction, display=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_placeholder(paragraph, text, **formatting):
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Aptos")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Aptos")
    if "bold" in formatting:
        run.bold = formatting["bold"]
    if "italic" in formatting:
        run.italic = formatting["italic"]
    if "font_size" in formatting:
        run.font.size = formatting["font_size"]
    if "font_color" in formatting:
        run.font.color.rgb = formatting["font_color"]
    return run


def add_label_value(doc, label, placeholder):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    add_placeholder(p, placeholder)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, before, after in (
        ("Heading 1", 17, 18, 9),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(RED if name != "Heading 3" else DARK)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_header_footer(section):
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    hp = section.header.paragraphs[0]
    hp.text = "HUAWEI CLOUD  |  COMMERCIAL PROPOSAL"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.runs[0]
    hr.font.name = "Aptos"
    hr.font.size = Pt(8)
    hr.font.bold = True
    hr.font.color.rgb = RGBColor.from_string(MUTED)
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("CONFIDENTIAL  •  ")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string(MUTED)
    add_field(fp, " PAGE ", "1")
    fp.add_run("  •  ${customer_name}")
    for run in fp.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)
configure_styles(doc)
configure_header_footer(section)

# Cover: proposal_centerpiece pattern.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(80)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HUAWEI CLOUD")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(RED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(8)
add_placeholder(p, "${proposal_title}", bold=True, font_size=Pt(28),
                font_color=RGBColor.from_string(DARK))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(45)
add_placeholder(p, "Prepared for ${customer_name}", italic=True,
                font_size=Pt(15), font_color=RGBColor.from_string(MUTED))

meta = doc.add_table(rows=4, cols=2)
meta.style = "Table Grid"
set_table_geometry(meta, [4680, 4680])
metadata = [
    ("Proposal date", "${proposal_date}", "Prepared by", "${prepared_by}"),
    ("Customer industry", "${customer_industry}", "Currency", "${currency}"),
    ("Validity", "${validity_period}", "Payment terms", "${payment_terms}"),
    ("Monthly total", "${quote_total_monthly}", "Annual total", "${quote_total_annual}"),
]
for row, values in zip(meta.rows, metadata):
    left, right = row.cells
    for cell, label, value in ((left, values[0], values[1]), (right, values[2], values[3])):
        shade(cell, LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        lr = p.add_run(label.upper())
        lr.bold = True
        lr.font.size = Pt(8)
        lr.font.color.rgb = RGBColor.from_string(MUTED)
        vp = cell.add_paragraph()
        vp.paragraph_format.space_after = Pt(0)
        add_placeholder(vp, value, bold=True)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CONFIDENTIAL")
r.bold = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor.from_string(RED)

doc.add_page_break()
doc.add_heading("Table of Contents", level=1)
p = doc.add_paragraph()
add_field(p, ' TOC \\o "1-3" \\h \\z \\u ', "Right-click and select Update Field in Microsoft Word.")
note = doc.add_paragraph("After opening the generated proposal, right-click this table of contents and choose Update Field > Update entire table.")
note.runs[0].italic = True
note.runs[0].font.size = Pt(9)
note.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

doc.add_page_break()
doc.add_heading("1. Executive Summary", level=1)
add_placeholder(doc.add_paragraph(), "${executive_summary}")

doc.add_heading("2. Scope of Work", level=1)
add_placeholder(doc.add_paragraph(), "${scope_of_work}")

doc.add_heading("3. Out of Scope", level=1)
add_placeholder(doc.add_paragraph(), "${out_of_scope}")

doc.add_heading("4. Assumptions", level=1)
add_placeholder(doc.add_paragraph(), "${assumptions}")

doc.add_heading("5. Commercial Summary", level=1)
p = doc.add_paragraph("The following pricing summary is based on the uploaded Huawei Cloud quote. Repeated services are consolidated and discounted prices are summed.")
p.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

summary = doc.add_table(rows=2, cols=7)
summary.style = "Table Grid"
widths = [2200, 1200, 1150, 650, 1250, 1460, 1450]
set_table_geometry(summary, widths)
headers = ["Service", "Region", "Billing", "Qty", "Average unit", "Discounted / month", "Annual"]
for cell, text in zip(summary.rows[0].cells, headers):
    shade(cell, RED)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(WHITE)
set_repeat_header(summary.rows[0])

row_placeholders = [
    "${quote_service_name}", "${quote_region}", "${quote_billing_mode}",
    "${quote_quantity}", "${quote_unit_price}", "${quote_discounted_price}",
    "${quote_annual_total}",
]
for cell, text in zip(summary.rows[1].cells, row_placeholders):
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if text != "${quote_service_name}":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_placeholder(p, text, font_size=Pt(8))

totals = doc.add_table(rows=2, cols=2)
totals.style = "Table Grid"
set_table_geometry(totals, [6500, 2860])
for row, label, value in (
    (totals.rows[0], "TOTAL MONTHLY AMOUNT", "${quote_total_monthly}"),
    (totals.rows[1], "TOTAL ANNUAL AMOUNT", "${quote_total_annual}"),
):
    shade(row.cells[0], LIGHT)
    shade(row.cells[1], LIGHT)
    lr = row.cells[0].paragraphs[0].add_run(label)
    lr.bold = True
    lr.font.size = Pt(9)
    vp = row.cells[1].paragraphs[0]
    vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_placeholder(vp, value, bold=True, font_size=Pt(10),
                    font_color=RGBColor.from_string(RED))

doc.add_heading("6. Huawei Cloud Services", level=1)
marker = doc.add_paragraph()
add_placeholder(marker, "${services_block}", font_size=Pt(1),
                font_color=RGBColor.from_string(WHITE))
heading = doc.add_paragraph(style="Heading 2")
add_placeholder(heading, "${service_heading}")
add_placeholder(doc.add_paragraph(), "${service_short_description}", italic=True)

p = doc.add_paragraph(style="Heading 3")
add_placeholder(p, "${service_definition_label}")
add_placeholder(doc.add_paragraph(), "${service_definition}")
p = doc.add_paragraph(style="Heading 3")
add_placeholder(p, "${service_details_label}")
add_placeholder(doc.add_paragraph(), "${service_details}")
p = doc.add_paragraph(style="Heading 3")
add_placeholder(p, "${service_key_benefits_label}")
add_placeholder(doc.add_paragraph(), "${service_key_benefits}")
p = doc.add_paragraph(style="Heading 3")
add_placeholder(p, "${service_typical_use_cases_label}")
add_placeholder(doc.add_paragraph(), "${service_typical_use_cases}")
p = doc.add_paragraph(style="Heading 3")
add_placeholder(p, "${service_proposal_positioning_label}")
add_placeholder(doc.add_paragraph(), "${service_proposal_positioning}")
p = doc.add_paragraph()
add_placeholder(p, "${service_official_link}")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_placeholder(p, "${service_diagram}")
marker = doc.add_paragraph()
add_placeholder(marker, "${/services_block}", font_size=Pt(1),
                font_color=RGBColor.from_string(WHITE))

doc.add_heading("7. Additional Notes", level=1)
add_placeholder(doc.add_paragraph(), "${additional_notes}")

doc.add_heading("8. Commercial Terms", level=1)
add_label_value(doc, "Currency", "${currency}")
add_label_value(doc, "Proposal validity", "${validity_period}")
add_label_value(doc, "Payment terms", "${payment_terms}")

doc.add_heading("9. Acceptance", level=1)
doc.add_paragraph("This section may be replaced with your organization’s approved signature and acceptance wording.")
signature = doc.add_table(rows=3, cols=2)
signature.style = "Table Grid"
set_table_geometry(signature, [4680, 4680])
for row, values in zip(signature.rows, [
    ("For ${customer_name}", "For the Service Provider"),
    ("Name: __________________________", "Name: __________________________"),
    ("Date: ___________________________", "Date: ___________________________"),
]):
    for cell, value in zip(row.cells, values):
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(10)
        add_placeholder(p, value)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
